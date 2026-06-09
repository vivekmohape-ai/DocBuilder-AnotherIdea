import streamlit as st
import zipfile
import io
import re
from pathlib import Path
from datetime import datetime
from PIL import Image as PILImage
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import base64

# ── Logo (embedded permanently) ──────────────────────────────────────────────
LOGO_B64 = "/9j/4AAQSkZJRgABAQAAAQABAAD/4gHYSUNDX1BST0ZJTEUAAQEAAAHIAAAAAAQwAABtbnRyUkdCIFhZWiAH4AABAAEAAAAAAABhY3NwAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAQAA9tYAAQAAAADTLQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAlkZXNjAAAA8AAAACRyWFlaAAABFAAAABRnWFlaAAABKAAAABRiWFlaAAABPAAAABR3dHB0AAABUAAAABRyVFJDAAABZAAAAChnVFJDAAABZAAAAChiVFJDAAABZAAAAChjcHJ0AAABjAAAADxtbHVjAAAAAAAAAAEAAAAMZW5VUwAAAAgAAAAcAHMAUgBHAEJYWVogAAAAAAAAb6IAADj1AAADkFhZWiAAAAAAAABimQAAt4UAABjaWFlaIAAAAAAAACSgAAAPhAAAts9YWVogAAAAAAAA9tYAAQAAAADTLXBhcmEAAAAAAAQAAAACZmYAAPKnAAANWQAAE9AAAApbAAAAAAAAAABtbHVjAAAAAAAAAAEAAAAMZW5VUwAAACAAAAAcAEcAbwBvAGcAbABlACAASQBuAGMALgAgADIAMAAxADb/2wBDAAUDBAQEAwUEBAQFBQUGBwwIBwcHBw8LCwkMEQ8SEhEPERETFhwXExQaFRERGCEYGh0dHx8fExciJCIeJBweHx7/2wBDAQUFBQcGBw4ICA4eFBEUHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh7/wAARCAAXAPMDASIAAhEBAxEB/8QAGwAAAgMBAQEAAAAAAAAAAAAAAAgFBgcEAQP/xABJEAABAgUCBAIFBgcPBQAAAAABAgMABAUGEQcSCBMhMRRBIjJRYXEJFhcjgaEVJDdCQ4KRNjg5UlNUV2JydZKztNHTdISUpLH/xAAUAQEAAAAAAAAAAAAAAAAAAAAA/8QAFBEBAAAAAAAAAAAAAAAAAAAAAP/aAAwDAQACEQMRAD8AbC/LjkrRs2rXLUFYl6dKrfUM4KiB0SPeTgD4wgJvXXZdjL1Y+edZFH/DXhiwJpYbCj6fRHbl59DH2Rtnyhl8iQtalWHKPYeqTni50JPUMtn0En+0vr+pFYa1f0SHDn9FZXVwTTOUXfA9PFevzO/8p1+EA2+n1yyl4WTR7nkSksVGVQ+ADnaoj0k/YrI+yE/4htYb70+4nppVKrM6/SZJMutVJW8rwziC0CsFPYZyTnyPWLf8nlfPjrYq1hTb2XqavxkkCe7KzhaR7gvB/XiEuSUlZ/5QdmRnpdqZlZhgNusupCkrSZNQIIPcQDRaYXxQ9QrOk7loMwHGH04daKgVsOD1m1jyI+/ofOMRvO7LmluN63bYl69UWqI/INrdkEvqDC1FDpJKOxPQfsii1ynXBwraqiu0ZuaqGnVbdSmaaxu5PU+gT5LSMlKvzh0Md1VrtKuXjlsqu0ScbnKfO0ll1h5B6KBQ9+wg5BHkRAWLid1pvCVv2U0p0uQRXZgoTMzSEhTiVr9VtG7ok46lR7Z6YiEl+HnXWelFT9T1jmWaksBQZE7MrTn2FYIxj3AxF2Qppv5Qys/hEtcwvTAZyOm4yw249+2HOgE2szVPVbR3ViSsTVufNWo86pKW511XMKUqOA6h3G5QB6FKuo90X3UnSbXOuX3V6vberLlIpE0/vlJITb6eSjaBtwkYHUHtFE+UV8Oqu2MhvJnfryNvfZub+3v2+2G6oAdFCp4fCw8JZsOb/W3bRnPvzAIbSWtc6jrlOaTt6s1VFSld+6aVPPck7Wg4cDv2OI3/AEh0t1ptu/JKr3dqi5XaOylYekjNPLDhKSE9FDHQ9Yzmyf4ROuf9x/pEw40ArWl943VPcad121OXDUpiiy/ieTIOTClMt7QjGE5wMZMNLCCT0pqHO8Yd5saZVCXkK74iYPNf2beVtTuHppUM9vKNOFucZOf3Y0j/ANb/AIoDo4ebwuuscUN80OqXDU52linP8PKPzClNNYdAG1JOBgRRKjNauX5xKXdZNr6j1SjNSsy+4yhc46Gm0IIG1IT27x08GTdaZ4krvauN5D9YRLPCdcRjCneYNxGAB39gitJlNRp7i2vdjTCoy8hXPEzJU6/s28rKdw9NKhnt5QGn/QhxHf03vf8AmzP+0XfVA33Y/CjUU1a5nJi5qewlJqso8tLixzQArccK3bTgmKZLW5xiCYaLt4UkthY3j8X7Z6/oo0ni1DyeG24hMEKeEsyHCPNXMRn74DC9KrG171DsOQuymazVCVZnCsBmYm39ydqinuM57R87tuDXjh6uCkVO6rs+dlBn3ihxtx5ToXtA3IysbkKwcgg4McOg/E7RdPNMabak5atSnnZQuFT7LqQhW5RV0z8Y4tTL9u7iirNJtKzrRfkqdIzBdfeW5zAkq9EOOqACUJAzgdST+yAeeh1GXrFGkqrKEmXnGEPtk/xVAEf/AGF+48bquS1bGoMzbVdqFIfen1occlH1NqWkIzgkeUb3a9KRQ7bptGbWXEyUq2wFH87akDP3QtfyjX5Prb/vJf8AlwHvFXeN10DQSyKtRbhqVPqE2WfETEvMKQ47lnJ3EdT16xAWVpfxDXTaNJuSU1ommGKnKNzTbbk7MbkJWkEA46Z6x7xk/vb9Pvix/p4Yrh+/IfZf9yy3+WIBZm9SdatBL6ptK1SnzcVuTpwJgr5pKMgKW25gK3JyCUq7j45ib46L+uegTFoTFoXRUqZLT8q66TJTKm0ujKdpOD16GJP5RlcqnTi3UOcvxKqmotZHpbQ2d2Pd1Tn7Ix/iuTNp0x0dTPb/ABQt5PN3HJ3bG85gNUY0V4jXmG3U63PALSFAGdmPMZ9kSFB0T18brck5WNapx2nJfQqaQxOzHMW2D6QTnpkjIiOlbc4xDLNKavGkhsoTsH4v0GOn6KNSsGb1FsbTGsVzVq4GKrVw4TKtNIbCGxjCEgoSncVKOT7BAQ2t14V927WbYtKcnm1U9hTk2ZVwhS1bdxBI6nagE/EmL/ojdS7psllyadLlQkz4eaKu6iOqVn25TjJ9uYqfDnb0w8Kje9WBcm6g6tLCljqUlWXF/rK6D3D3xDUY/Rfrc5SipSaNWSCjPqpC1HZ/hVlPwMBv0EEEAnNrWZUNXuLCq3RdlMbctqnbzLy7ziFpdab+rZQUZJGSd5BGM5EMf9EWl+7d8wLcz/0Df+0EEAuN12bP6TcWdJua0KU03b1Q5ZflmFobS2hz6t1ASSOnTeOnc+6JiqW7WFce8nXEygNPCE5d5qf5qR6uc9/dBBAMteNt0a7bbnber8kicp843sdbV9ygfJQPUEdiITXTLRq4tNeK6iSjgROUht5x6Vmw6nKmShYTuTnIUOx6Y8xBBAaRxQ6IXNWLxlNUdNZtDFxSexb7BdS2Vqb9V1ClejuA6EK6EAfA5hM8VWslEnmbfqlBtpypgJQVLaVuWT0BJQ7sz8MCCCAsOmuk2pGrOpUjqRqw9LN0uXWlxmVQ6hXNSg5S2hCCQhGe+Tkw48EEAp1n27WGePasVxyUAp7gfKXeak5zKpHq5z390NjBBAKlpXbtYleN67ay/KBMi74nY7zEnOUox0Bz90NbBBAKlw529WKfxUX3U5uUDco/z+WvmJO7LoI6A5jOq9O6iafcSt4Xha1Cp88t+bfaSJx1JQUKIOcBaTnoIIIC0jiJ17z1sq1f8Sv+eL3etXvPULhDrk/cNMkJatzSilMtJq2thCXk46qWoZwCe8EEBb+EqjPU/QOi06qybSXhzg42rasEFauhxkHpGI3laFzaA68MXbYcqJu2asSqYp4fQjDZOXGsKI7esg+XaCCAcSlTrVSpktUGAoNTLSXUhQ6gKGcH3wufH7Q6pXLEt9mly3iFt1BaljmJTgbP6xEEEBD8WtuViq8P9i0+RlA7MS5Z5qOalO36jHckA9Yptm6164WvalLtyQs+2nJWmyrcqyt1ZK1JQnAKsPAZ6eyCCAkqdpJrBrde9PuHVmYlafQZQhSZdl5tW5skKKG0Nk43dAVKOce2O/j4tSqVWds9ihyCXGJWWeb2hxCAgApwPSI8hBBATGkusOt9z3vRrcnbYteSkHXAJmYG4qQykZUQA8fSwMDoepi+65M1e7LupVoU9vbJtOJW+4VpA5ix3wTk7UEn9YwQQGv0iQlqVSpWmyTYblpVpLTSR5JSMCM/4gbUXcFqIqEkAKhTF8xByAVtnotOT28j9nvgggOmzr+l/mxIJrqXkVJDXLmMDcFKSSndkdOoAP2wQQQH/9k="

def logo_bytes():
    return base64.b64decode(LOGO_B64)

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="DocBuilder — Another Idea",
    page_icon="🏙️",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── CSS ───────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@400;500;600;700;800&family=Instrument+Sans:wght@300;400;500&display=swap');

:root {
  --bg: #0A0A0F;
  --surface: #13131A;
  --surface2: #1C1C28;
  --border: #2A2A3A;
  --border-bright: #3D3D55;
  --accent: #6C63FF;
  --accent2: #FF5C87;
  --accent-glow: rgba(108,99,255,0.25);
  --text: #F0F0F8;
  --muted: #7070A0;
  --success: #00E5A0;
  --success-bg: rgba(0,229,160,0.08);
  --tag-bg: rgba(108,99,255,0.1);
}

html, body, [class*="css"] {
  font-family: 'Instrument Sans', sans-serif !important;
  background-color: var(--bg) !important;
  color: var(--text) !important;
}
.main { background-color: var(--bg) !important; }
.block-container { padding: 0 2.5rem 4rem 2.5rem !important; max-width: 1080px; }

/* ── Hero ── */
.hero {
  padding: 3.5rem 0 3rem 0;
  text-align: center;
  position: relative;
}
.hero::after {
  content: '';
  display: block;
  height: 1px;
  background: linear-gradient(90deg, transparent, var(--border-bright), transparent);
  margin-top: 2.5rem;
}
.hero-eyebrow {
  font-family: 'Syne', sans-serif;
  font-size: 0.65rem;
  font-weight: 600;
  letter-spacing: 0.35em;
  text-transform: uppercase;
  color: var(--accent);
  margin: 0 0 1rem 0;
}
.hero-title {
  font-family: 'Syne', sans-serif;
  font-size: 3.8rem;
  font-weight: 800;
  line-height: 1;
  letter-spacing: -0.02em;
  margin: 0;
  background: linear-gradient(135deg, #fff 30%, var(--accent) 100%);
  -webkit-background-clip: text;
  -webkit-text-fill-color: transparent;
  background-clip: text;
}
.hero-sub {
  font-size: 0.88rem;
  color: var(--muted);
  margin: 0.9rem 0 0 0;
  letter-spacing: 0.02em;
}

/* ── Cards ── */
.card {
  background: var(--surface);
  border: 1px solid var(--border);
  border-radius: 16px;
  padding: 1.6rem;
  margin-bottom: 1.2rem;
  transition: border-color 0.2s;
}
.card:hover { border-color: var(--border-bright); }

/* ── Step label ── */
.step-badge {
  display: inline-flex;
  align-items: center;
  gap: 0.45rem;
  background: var(--tag-bg);
  border: 1px solid rgba(108,99,255,0.3);
  border-radius: 999px;
  padding: 0.22rem 0.75rem;
  font-family: 'Syne', sans-serif;
  font-size: 0.65rem;
  font-weight: 700;
  letter-spacing: 0.18em;
  text-transform: uppercase;
  color: var(--accent);
  margin-bottom: 0.7rem;
}
.step-heading {
  font-family: 'Syne', sans-serif;
  font-size: 1.15rem;
  font-weight: 700;
  color: var(--text);
  margin: 0 0 1rem 0;
  letter-spacing: -0.01em;
}

/* ── File uploader ── */
.stFileUploader > label { display: none !important; }
[data-testid="stFileUploader"] {
  border: 1.5px dashed var(--border-bright) !important;
  border-radius: 12px !important;
  background: var(--surface2) !important;
  transition: border-color 0.2s, background 0.2s !important;
}
[data-testid="stFileUploader"]:hover {
  border-color: var(--accent) !important;
  background: rgba(108,99,255,0.05) !important;
}
[data-testid="stFileUploader"] * { color: var(--muted) !important; }

/* ── Textarea ── */
.stTextArea textarea {
  font-family: 'Instrument Sans', monospace !important;
  font-size: 0.8rem !important;
  border: 1px solid var(--border) !important;
  border-radius: 12px !important;
  background: var(--surface2) !important;
  color: var(--text) !important;
  line-height: 1.8 !important;
  caret-color: var(--accent) !important;
  transition: border-color 0.2s, box-shadow 0.2s !important;
}
.stTextArea textarea:focus {
  border-color: var(--accent) !important;
  box-shadow: 0 0 0 3px var(--accent-glow) !important;
  outline: none !important;
}
.stTextArea textarea::placeholder { color: var(--muted) !important; opacity: 0.6 !important; }

/* ── Primary button ── */
.stButton > button {
  font-family: 'Syne', sans-serif !important;
  font-size: 0.78rem !important;
  font-weight: 700 !important;
  letter-spacing: 0.08em !important;
  text-transform: uppercase !important;
  background: linear-gradient(135deg, var(--accent), #9B93FF) !important;
  color: #fff !important;
  border: none !important;
  border-radius: 999px !important;
  padding: 0.72rem 2.2rem !important;
  box-shadow: 0 4px 20px var(--accent-glow) !important;
  transition: transform 0.18s ease, box-shadow 0.18s ease, filter 0.18s ease !important;
  cursor: pointer !important;
}
.stButton > button:hover {
  transform: translateY(-2px) !important;
  box-shadow: 0 8px 30px rgba(108,99,255,0.45) !important;
  filter: brightness(1.1) !important;
}
.stButton > button:active {
  transform: translateY(0px) !important;
  box-shadow: 0 2px 10px var(--accent-glow) !important;
}

/* ── Download button ── */
[data-testid="stDownloadButton"] > button {
  font-family: 'Syne', sans-serif !important;
  font-size: 0.78rem !important;
  font-weight: 700 !important;
  letter-spacing: 0.08em !important;
  text-transform: uppercase !important;
  background: linear-gradient(135deg, var(--success), #00C488) !important;
  color: #0A0A0F !important;
  border: none !important;
  border-radius: 999px !important;
  padding: 0.72rem 2.2rem !important;
  width: 100% !important;
  box-shadow: 0 4px 20px rgba(0,229,160,0.25) !important;
  transition: transform 0.18s ease, box-shadow 0.18s ease, filter 0.18s ease !important;
}
[data-testid="stDownloadButton"] > button:hover {
  transform: translateY(-2px) !important;
  box-shadow: 0 8px 30px rgba(0,229,160,0.4) !important;
  filter: brightness(1.08) !important;
}

/* ── Metrics ── */
.metric-row { display: flex; gap: 0.8rem; margin: 1rem 0; }
.metric-card {
  flex: 1;
  background: var(--surface2);
  border: 1px solid var(--border);
  border-radius: 12px;
  padding: 1rem;
  text-align: center;
}
.metric-number {
  font-family: 'Syne', sans-serif;
  font-size: 2.2rem;
  font-weight: 800;
  background: linear-gradient(135deg, var(--accent), var(--accent2));
  -webkit-background-clip: text;
  -webkit-text-fill-color: transparent;
  background-clip: text;
  line-height: 1;
}
.metric-label {
  font-size: 0.65rem;
  letter-spacing: 0.15em;
  text-transform: uppercase;
  color: var(--muted);
  margin-top: 0.35rem;
}

/* ── Tags ── */
.preview-tag {
  display: inline-block;
  background: var(--tag-bg);
  border: 1px solid rgba(108,99,255,0.25);
  border-radius: 6px;
  padding: 0.25rem 0.6rem;
  font-size: 0.72rem;
  color: #A0A0D0;
  margin: 0.2rem 0.2rem 0.2rem 0;
  font-family: 'Instrument Sans', monospace;
}

/* ── Format hint ── */
.format-hint {
  background: var(--surface2);
  border: 1px solid var(--border);
  border-radius: 12px;
  padding: 1.2rem 1.4rem;
  font-size: 0.78rem;
  color: var(--muted);
  line-height: 2;
  margin-bottom: 0.8rem;
}
.format-hint strong { color: var(--text); font-weight: 500; }
.format-hint code {
  background: rgba(108,99,255,0.15);
  border: 1px solid rgba(108,99,255,0.2);
  padding: 0.1rem 0.4rem;
  border-radius: 5px;
  color: #B0AAFF;
  font-size: 0.74rem;
  font-family: 'Instrument Sans', monospace;
}

/* ── Success / status ── */
.success-box {
  background: var(--success-bg);
  border: 1px solid rgba(0,229,160,0.3);
  border-radius: 12px;
  padding: 1rem 1.4rem;
  color: var(--success);
  font-size: 0.85rem;
  margin: 0.8rem 0;
}

/* ── Cover preview ── */
.cover-preview {
  background: var(--surface2);
  border: 1px solid var(--border);
  border-left: 3px solid var(--accent);
  border-radius: 12px;
  padding: 1.2rem 1.4rem;
  margin-top: 1rem;
}
.cover-preview .label {
  font-size: 0.62rem;
  letter-spacing: 0.2em;
  text-transform: uppercase;
  color: var(--accent);
  margin: 0 0 0.6rem 0;
  font-family: 'Syne', sans-serif;
  font-weight: 700;
}
.cover-preview .title { font-size: 0.9rem; color: var(--text); font-weight: 500; margin: 0.2rem 0; }
.cover-preview .sub { font-size: 0.78rem; color: var(--muted); margin: 0; }

/* ── Divider ── */
.divider {
  height: 1px;
  background: linear-gradient(90deg, transparent, var(--border-bright), transparent);
  border: none;
  margin: 2rem 0;
}

/* ── Expander ── */
.streamlit-expanderHeader {
  font-family: 'Syne', sans-serif !important;
  font-size: 0.75rem !important;
  color: var(--muted) !important;
  letter-spacing: 0.06em !important;
  background: var(--surface2) !important;
  border-radius: 8px !important;
}
.streamlit-expanderContent {
  background: var(--surface2) !important;
  border-radius: 0 0 8px 8px !important;
}

/* ── Progress bar ── */
.stProgress > div > div {
  background: linear-gradient(90deg, var(--accent), var(--accent2)) !important;
  border-radius: 999px !important;
}
.stProgress > div {
  background: var(--surface2) !important;
  border-radius: 999px !important;
}

/* ── Spinner ── */
.stSpinner > div { border-top-color: var(--accent) !important; }

/* ── Alerts ── */
.stAlert { border-radius: 10px !important; }

/* ── Scrollbar ── */
::-webkit-scrollbar { width: 6px; }
::-webkit-scrollbar-track { background: var(--bg); }
::-webkit-scrollbar-thumb { background: var(--border-bright); border-radius: 3px; }

/* ── Hide streamlit chrome ── */
#MainMenu, footer, header { visibility: hidden; }
</style>
""", unsafe_allow_html=True)


# ── Helpers ───────────────────────────────────────────────────────────────────
def parse_flow_text(flow_text: str):
    sections = []
    current_project = None
    current_entries = []
    for raw_line in flow_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        project_match = (
            re.match(r'^\[(.+)\]$', line) or
            re.match(r'^PROJECT\s*:\s*(.+)', line, re.IGNORECASE) or
            re.match(r'^#{1,3}\s+(.+)', line)
        )
        if project_match:
            if current_project and current_entries:
                sections.append({"project": current_project, "entries": current_entries})
            current_project = project_match.group(1).strip()
            current_entries = []
            continue
        entry_match = re.match(r'^(.+?)\s*[|\-,]\s*(.+)$', line)
        if entry_match and current_project:
            current_entries.append({"filename": entry_match.group(1).strip(), "caption": entry_match.group(2).strip()})
        elif line and current_project:
            current_entries.append({"filename": line, "caption": ""})
    if current_project and current_entries:
        sections.append({"project": current_project, "entries": current_entries})
    return sections


def extract_images_from_zip(zip_bytes: bytes):
    image_map = {}
    folder_structure = {}
    img_exts = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.tiff'}
    with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
        for name in zf.namelist():
            p = Path(name)
            if p.suffix.lower() in img_exts and not name.startswith('__MACOSX'):
                try:
                    data = zf.read(name)
                    img = PILImage.open(io.BytesIO(data)).convert("RGB")
                    image_map[p.name.lower()] = img
                    folder = str(p.parent) if str(p.parent) != '.' else 'root'
                    folder_structure.setdefault(folder, []).append(p.name)
                except Exception:
                    pass
    return image_map, folder_structure


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for b in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{b}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def set_cell_width(cell, width_inches):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def add_image_to_cell(cell, img_bytes_io, width_inches):
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(img_bytes_io, width=Inches(width_inches))


def add_caption_to_cell(cell, caption_text):
    para = cell.add_paragraph(caption_text)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0] if para.runs else para.add_run(caption_text)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 95, 90)
    run.font.italic = True
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(6)


def set_section_footer(doc_section, folder_name, logo_b):
    """Set centered footer with folder name for a given Word section."""
    doc_section.footer_distance = Cm(1.0)
    footer = doc_section.footer
    footer.is_linked_to_previous = False

    # Clear existing paragraphs
    for p in footer.paragraphs:
        p.clear()

    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    run = fp.add_run(folder_name)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(138, 132, 128)
    run.font.name = "Calibri"


def set_section_header(doc_section, logo_b):
    """Set header with Another Idea logo for a given Word section."""
    doc_section.header_distance = Cm(0.8)
    header = doc_section.header
    header.is_linked_to_previous = False

    for p in header.paragraphs:
        p.clear()

    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(4)
    run = hp.add_run()
    run.add_picture(io.BytesIO(logo_b), width=Inches(1.2))


def add_section_break(doc):
    """Add a next-page section break via XML."""
    # Add a paragraph with a section break
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), '12240')   # A4-ish
    pgSz.set(qn('w:h'), '15840')
    sectPr.append(pgSz)
    pgMar = OxmlElement('w:pgMar')
    pgMar.set(qn('w:top'), '1134')
    pgMar.set(qn('w:right'), '1134')
    pgMar.set(qn('w:bottom'), '1134')
    pgMar.set(qn('w:left'), '1134')
    pgMar.set(qn('w:header'), '720')
    pgMar.set(qn('w:footer'), '720')
    sectPr.append(pgMar)
    # Next page type
    pgType = OxmlElement('w:type')
    pgType.set(qn('w:val'), 'nextPage')
    sectPr.insert(0, pgType)
    pPr.append(sectPr)
    return para


# ── Core doc builder ──────────────────────────────────────────────────────────
def build_word_doc(sections, image_map, zip_name, progress_cb=None):
    doc = Document()
    logo_b = logo_bytes()
    usable_width = 6.3

    # ── Global page setup ─────────────────────────────────────────────────────
    for sec in doc.sections:
        sec.page_width  = Cm(21)
        sec.page_height = Cm(29.7)
        sec.top_margin    = Cm(2.2)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(2.0)
        sec.right_margin  = Cm(2.0)
        sec.header_distance = Cm(0.8)
        sec.footer_distance = Cm(1.0)

    # ── COVER PAGE (section 0) ────────────────────────────────────────────────
    cover_section = doc.sections[0]

    # No header/footer on cover
    cover_section.header.is_linked_to_previous = False
    cover_section.footer.is_linked_to_previous = False
    for p in cover_section.header.paragraphs:
        p.clear()
    for p in cover_section.footer.paragraphs:
        p.clear()

    # Vertical spacer
    for _ in range(8):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    # Logo
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_para.paragraph_format.space_after = Pt(32)
    logo_run = logo_para.add_run()
    logo_run.add_picture(io.BytesIO(logo_b), width=Inches(2.2))

    # SUPPORTINGS FOR line
    sup_para = doc.add_paragraph()
    sup_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sup_para.paragraph_format.space_after = Pt(12)
    sup_run = sup_para.add_run("SUPPORTINGS FOR")
    sup_run.font.size = Pt(11)
    sup_run.font.color.rgb = RGBColor(138, 132, 128)
    sup_run.font.name = "Calibri"
    sup_run.font.bold = False

    # Project name (ZIP filename without extension)
    zip_display = Path(zip_name).stem.replace("_", " ").replace("-", " ").upper()
    proj_para = doc.add_paragraph()
    proj_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    proj_para.paragraph_format.space_after = Pt(20)
    proj_run = proj_para.add_run(f"({zip_display})")
    proj_run.font.size = Pt(16)
    proj_run.font.bold = True
    proj_run.font.color.rgb = RGBColor(28, 28, 28)
    proj_run.font.name = "Calibri"

    # Month Year
    month_year = datetime.now().strftime("%B %Y").upper()
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(month_year)
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(138, 132, 128)
    date_run.font.name = "Calibri"
    date_para.paragraph_format.space_after = Pt(0)

    # ── IMAGE SECTIONS (one Word section per folder) ──────────────────────────
    total_entries = sum(len(s['entries']) for s in sections)
    processed = 0

    for sec_idx, section_data in enumerate(sections):
        folder_name = section_data["project"]
        entries = section_data["entries"]

        # Insert a next-page section break paragraph (carries sectPr for PREVIOUS section)
        add_section_break(doc)

        # Now add content for this folder
        # Section heading
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run(folder_name.upper())
        heading_run.font.size = Pt(12)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(28, 28, 28)
        heading_run.font.name = "Calibri"
        heading_para.paragraph_format.space_before = Pt(6)
        heading_para.paragraph_format.space_after = Pt(3)

        rule_para = doc.add_paragraph()
        rule_run = rule_para.add_run("─" * 65)
        rule_run.font.size = Pt(7)
        rule_run.font.color.rgb = RGBColor(184, 147, 85)
        rule_para.paragraph_format.space_after = Pt(8)

        # Images
        i = 0
        while i < len(entries):
            entry = entries[i]
            fname = entry["filename"].lower()
            caption = entry["caption"]

            img = image_map.get(fname)
            if img is None:
                for key in image_map:
                    if fname in key or key in fname:
                        img = image_map[key]
                        break

            if img is None:
                mp = doc.add_paragraph(f"⚠ Image not found: {entry['filename']}")
                mp.runs[0].font.color.rgb = RGBColor(180, 80, 80)
                mp.runs[0].font.size = Pt(9)
                i += 1
                processed += 1
                if progress_cb: progress_cb(min(processed/total_entries, 1.0))
                continue

            w, h = img.size
            is_landscape = w >= h

            # Try to pair with next if both landscape
            paired = False
            if is_landscape and i + 1 < len(entries):
                next_entry = entries[i + 1]
                next_fname = next_entry["filename"].lower()
                next_img = image_map.get(next_fname)
                if next_img is None:
                    for key in image_map:
                        if next_fname in key or key in next_fname:
                            next_img = image_map[key]
                            break
                if next_img:
                    nw, nh = next_img.size
                    if nw >= nh:
                        paired = True
                        col_w = (usable_width - 0.2) / 2
                        table = doc.add_table(rows=2, cols=2)
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        remove_table_borders(table)
                        for ci, (cur_img, cur_entry) in enumerate([(img, entry), (next_img, next_entry)]):
                            set_cell_width(table.cell(0, ci), col_w)
                            set_cell_width(table.cell(1, ci), col_w)
                            img_io = io.BytesIO()
                            cur_img.save(img_io, format='JPEG', quality=92)
                            img_io.seek(0)
                            cell = table.cell(0, ci)
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            add_image_to_cell(cell, img_io, col_w - 0.08)
                            if cur_entry["caption"]:
                                add_caption_to_cell(table.cell(1, ci), cur_entry["caption"])
                        sp = doc.add_paragraph()
                        sp.paragraph_format.space_after = Pt(8)
                        i += 2
                        processed += 2
                        if progress_cb: progress_cb(min(processed/total_entries, 1.0))
                        continue

            # Single image
            img_io = io.BytesIO()
            img.save(img_io, format='JPEG', quality=92)
            img_io.seek(0)
            img_width = usable_width if is_landscape else min(3.6, usable_width * 0.55)
            table = doc.add_table(rows=2, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            remove_table_borders(table)
            set_cell_width(table.cell(0, 0), usable_width)
            set_cell_width(table.cell(1, 0), usable_width)
            img_cell = table.cell(0, 0)
            img_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            add_image_to_cell(img_cell, img_io, img_width)
            if caption:
                add_caption_to_cell(table.cell(1, 0), caption)
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(8)
            i += 1
            processed += 1
            if progress_cb: progress_cb(min(processed/total_entries, 1.0))

    # ── Apply headers & footers to each Word section ──────────────────────────
    # After building content, doc.sections has:
    #   [0] = cover,  [1..N] = one per folder
    # The section breaks we inserted via pPr.sectPr define sections 0..N-1.
    # The final doc.sections[-1] is the last body section.

    # Re-apply page margins to all sections
    for sec in doc.sections:
        sec.page_width  = Cm(21)
        sec.page_height = Cm(29.7)
        sec.top_margin    = Cm(2.2)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(2.0)
        sec.right_margin  = Cm(2.0)
        sec.header_distance = Cm(0.8)
        sec.footer_distance = Cm(1.0)

    # Section 0 = cover (no header/footer — already cleared)
    # Sections 1..N = image sections, assign folder name footers
    content_sections = doc.sections[1:]  # skip cover
    for idx, ws in enumerate(content_sections):
        fname = sections[idx]["project"] if idx < len(sections) else ""
        set_section_header(ws, logo_b)
        set_section_footer(ws, fname, logo_b)

    out_io = io.BytesIO()
    doc.save(out_io)
    out_io.seek(0)
    return out_io


# ── UI ────────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="hero">
    <p class="hero-eyebrow">Another Idea · Real Estate</p>
    <h1 class="hero-title">DocBuilder</h1>
    <p class="hero-sub">Upload images · define sequence · download a formatted Word doc</p>
</div>
""", unsafe_allow_html=True)

col1, col2 = st.columns([1, 1], gap="large")

# ── SESSION STATE ─────────────────────────────────────────────────────────────
if "image_map" not in st.session_state:
    st.session_state.image_map = {}
if "folder_structure" not in st.session_state:
    st.session_state.folder_structure = {}
if "zip_name" not in st.session_state:
    st.session_state.zip_name = ""

with col1:
    st.markdown('<div class="step-badge">⬆ Step 01</div>', unsafe_allow_html=True)
    st.markdown('<p class="step-heading">Upload Image Archive</p>', unsafe_allow_html=True)

    st.markdown('<div class="card">', unsafe_allow_html=True)
    zip_file = st.file_uploader("Upload ZIP", type=["zip"])

    if zip_file:
        if zip_file.name != st.session_state.zip_name:
            with st.spinner("Scanning archive..."):
                zip_bytes = zip_file.read()
                st.session_state.image_map, st.session_state.folder_structure = extract_images_from_zip(zip_bytes)
                st.session_state.zip_name = zip_file.name

        total_imgs = len(st.session_state.image_map)
        total_folders = len(st.session_state.folder_structure)
        st.markdown(f"""
        <div class="metric-row">
            <div class="metric-card">
                <div class="metric-number">{total_imgs}</div>
                <div class="metric-label">Images</div>
            </div>
            <div class="metric-card">
                <div class="metric-number">{total_folders}</div>
                <div class="metric-label">Folders</div>
            </div>
        </div>""", unsafe_allow_html=True)
        with st.expander("View folders & files"):
            for folder, files in sorted(st.session_state.folder_structure.items()):
                st.markdown(f"<span style='color:#A0A0D0;font-size:0.8rem;font-weight:500'>📁 {folder}</span> <span style='color:var(--muted);font-size:0.72rem'>({len(files)} images)</span>", unsafe_allow_html=True)
                for f in files[:6]:
                    st.markdown(f'<span class="preview-tag">{f}</span>', unsafe_allow_html=True)
                if len(files) > 6:
                    st.markdown(f"<span style='color:var(--muted);font-size:0.72rem'>+{len(files)-6} more</span>", unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    if st.session_state.zip_name:
        zip_display = Path(st.session_state.zip_name).stem.replace("_", " ").replace("-", " ").upper()
        month_year = datetime.now().strftime("%B %Y").upper()
        st.markdown(f"""
        <div class="cover-preview">
            <p class="label">Cover Page Preview</p>
            <p class="title">ANOTHER IDEA <span style="opacity:0.4">[logo]</span></p>
            <p class="title">SUPPORTINGS FOR ({zip_display})</p>
            <p class="sub">{month_year}</p>
        </div>
        """, unsafe_allow_html=True)

with col2:
    st.markdown('<div class="step-badge">✏ Step 02</div>', unsafe_allow_html=True)
    st.markdown('<p class="step-heading">Image Sequence & Captions</p>', unsafe_allow_html=True)

    st.markdown("""
    <div class="format-hint">
    <strong>Format</strong> — folder name in brackets, then <code>filename | caption</code> per line<br><br>
    <code>[Empress Hill 4 BHK]</code><br>
    <code>screen1.jpg | Digital Screen — Bedroom</code><br>
    <code>screen2.jpg | Digital Screen — Living Room</code><br>
    <code>ooh.jpg | OOH Billboard</code><br><br>
    <code>[Safety Week 2026]</code><br>
    <code>badge.png | Safety Week Badge</code><br>
    <code>banner.jpg | Safety Week Banner</code>
    </div>
    """, unsafe_allow_html=True)

    flow_text = st.text_area(
        "Sequence",
        height=260,
        placeholder="[Folder Name]\nfilename.jpg | Caption text\nfilename2.jpg | Caption text\n\n[Next Folder]\n...",
        label_visibility="collapsed",
        key="flow_text_input"
    )

    if flow_text.strip():
        parsed = parse_flow_text(flow_text)
        if parsed:
            total_e = sum(len(s['entries']) for s in parsed)
            st.markdown(f"""
            <div class="success-box">
                ✓ &nbsp;<strong>{len(parsed)} sections</strong> &nbsp;·&nbsp; <strong>{total_e} images</strong> ready to build
            </div>""", unsafe_allow_html=True)
            with st.expander("Preview sections"):
                for s in parsed:
                    st.markdown(f"<span style='color:var(--accent);font-size:0.8rem;font-weight:600'>{s['project']}</span> <span style='color:var(--muted);font-size:0.75rem'>— {len(s['entries'])} images</span>", unsafe_allow_html=True)
                    for e in s['entries'][:3]:
                        st.markdown(f"<span style='color:var(--muted);font-size:0.73rem;padding-left:1rem'>→ {e['filename']} · {e['caption'] or '(no caption)'}</span>", unsafe_allow_html=True)
                    if len(s['entries']) > 3:
                        st.markdown(f"<span style='color:var(--muted);font-size:0.72rem;padding-left:1rem'>+{len(s['entries'])-3} more</span>", unsafe_allow_html=True)

# ── Generate ──────────────────────────────────────────────────────────────────
st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

gc1, gc2, gc3 = st.columns([1.8, 1, 1.8])
with gc2:
    generate = st.button("Generate Doc", use_container_width=True)

# Read from session_state so values survive the rerun triggered by the button
flow_text = st.session_state.get("flow_text_input", "")
image_map = st.session_state.image_map
zip_name  = st.session_state.zip_name

if generate:
    if not st.session_state.zip_name or not st.session_state.image_map:
        st.error("Please upload a ZIP file first.")
    elif not flow_text.strip():
        st.error("Please paste your image sequence.")
    else:
        parsed_sections = parse_flow_text(flow_text)
        if not parsed_sections:
            st.error("Could not parse flow text — check the format above.")
        else:
            status_text = st.empty()
            progress_bar = st.progress(0)
            status_text.markdown("<span style='color:var(--muted);font-size:0.8rem'>Building your document…</span>", unsafe_allow_html=True)

            def update_progress(val):
                progress_bar.progress(val)

            try:
                doc_io = build_word_doc(
                    sections=parsed_sections,
                    image_map=image_map,
                    zip_name=zip_name,
                    progress_cb=update_progress
                )
                progress_bar.progress(1.0)
                status_text.empty()
                st.markdown('<div class="success-box" style="text-align:center;font-size:0.9rem">✓ &nbsp; Document ready</div>', unsafe_allow_html=True)
                dl1, dl2, dl3 = st.columns([1.5, 1, 1.5])
                with dl2:
                    safe_name = Path(zip_name).stem
                    st.download_button(
                        label="↓ Download .docx",
                        data=doc_io,
                        file_name=f"{safe_name}_Supportings.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            except Exception as e:
                st.error(f"Error: {str(e)}")
                st.exception(e)

st.markdown("""
<div style='text-align:center;padding:3rem 0 1.5rem 0;border-top:1px solid var(--border);margin-top:3rem'>
    <p style='font-family:Syne,sans-serif;font-size:0.65rem;letter-spacing:0.2em;text-transform:uppercase;color:var(--muted)'>
        DocBuilder &nbsp;·&nbsp; Another Idea &nbsp;·&nbsp; Real Estate Marketing
    </p>
</div>
""", unsafe_allow_html=True)
